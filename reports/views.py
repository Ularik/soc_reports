import json
from django.shortcuts import render
from datetime import datetime, timedelta, time
from io import BytesIO
from django.db.models import Q
from django.http import FileResponse, JsonResponse
from django.contrib import messages
from django.http import HttpResponse
from django.shortcuts import redirect, get_object_or_404
from django.urls import reverse_lazy
from django.views import View
from django.views.generic import CreateView, TemplateView, ListView, DetailView
from django.db.models import Count, Min, Max
from docx import Document
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from .utils import create_stat_report
from .models import Report, Pattern, Organization, RISK_LEVELS
from .forms import ReportsForm
from django.forms.models import model_to_dict
import base64
from django.contrib.auth import get_user_model
import requests
from django.contrib.auth.decorators import login_required
from django.core.cache import cache


User = get_user_model()

# Регистрация TTF-шрифта (путь укажите свой)
try:
    pdfmetrics.registerFont(
        TTFont('TimesNewRoman', r'/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf')
    )
except:
    pdfmetrics.registerFont(
        TTFont('TimesNewRoman', r'C:\Windows\Fonts\times.ttf')
    )

@login_required(login_url='user/login/')
def report_create_view(request):
    template_name = 'reports/report_waf_form.html'

    if request.method == 'POST':
        form = ReportsForm(request.POST)
        success_url = reverse_lazy('reports:report_new')

        if form.is_valid():
            # Получаем объект формы, но не сохраняем его сразу в БД
            report = form.save(commit=False)
            report.user = request.user
            # Сохраняем объект в БД
            report.save()

            # 2) Добавляем уведомление
            messages.success(request, 'Отчет отправлен')

            # 3) Редиректим на чистую форму с параметром download=<pk>
            return redirect(f"{success_url}?download={report.pk}")

        messages.success(request, 'Отчет некорректен')

    form = ReportsForm()
    # JSON для автозаполнения описания по типу атаки
    attack_types = {}
    for obj in Pattern.objects.all():
        data = model_to_dict(obj)
        attack_types[obj.pk] = data
    return render(request, template_name, context={
        'pattern': json.dumps(attack_types, ensure_ascii=False),
        'form': form
        })


class ReportDownloadView(View):
    def get(self, request, pk):
        report = get_object_or_404(Report, pk=pk)
        buffer = BytesIO()

        try:
            pdfmetrics.registerFont(TTFont('TimesNewRoman', r'/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf'))
        except:
            pdfmetrics.registerFont(TTFont('TimesNewRoman', r'C:\Windows\Fonts\times.ttf'))

        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            rightMargin=40, leftMargin=40,
            topMargin=60, bottomMargin=40
        )
        styles = getSampleStyleSheet()
        styles['Title'].fontName = 'TimesNewRoman'
        styles['Normal'].fontName = 'TimesNewRoman'
        cell = ParagraphStyle('CellStyle', fontName='TimesNewRoman', fontSize=10, leading=12, wordWrap='LTR')

        elems = [Paragraph("Отчет по выявлению аномалий средствами SOC", styles['Title']), Spacer(1, 12)]
        data = [
            [Paragraph("Поле", cell), Paragraph("Значение", cell)],
            [Paragraph("Дата и время", cell), Paragraph(report.detection_date.strftime('%Y-%m-%d'), cell)],
            [Paragraph("Тип угрозы", cell), Paragraph(report.attack_type, cell)],
            [Paragraph("Источник угрозы", cell), Paragraph(report.source_ip, cell)],

            [Paragraph("Адрес назначения", cell), Paragraph(report.destination_ip, cell)],
            [Paragraph("Средство обнаружения", cell), Paragraph(report.detection_tool, cell)],
            [Paragraph("Краткое описание", cell), Paragraph(report.short_description, cell)],
            [Paragraph("Методы атаки", cell), Paragraph(report.methods, cell)],
            [Paragraph("Протоколы и порты", cell), Paragraph(report.protocols_ports, cell)],
            [Paragraph("Критичность", cell), Paragraph(report.risk_assessment, cell)],
            [Paragraph("Потенциальные последствия", cell), Paragraph(report.potential_impact, cell)],
            [Paragraph(["Payload", "Data"][bool(report.host)], cell), Paragraph(report.data_or_payload, cell)],
            [Paragraph("Реагирование на ицнидент", cell),
             Paragraph(report.response_actions.replace('\n', '<br/>'), cell)],
        ]
        if report.detection_tool == 'WAF' and report.host:
            data.insert(4, [Paragraph("Host", cell), Paragraph(report.host, cell)])
        elif report.detection_tool == 'IPS' and report.cve:
            data.insert(4, [Paragraph("CVE", cell), Paragraph(report.cve, cell)])

        table = Table(data, colWidths=[150, 330])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elems.append(table)
        doc.build(elems)
        buffer.seek(0)

        name = f'{report.organization.name_en} {report.source_ip}.pdf'

        # Добавляем по АПИ отчет в cert.gov
        files = {
            'file': (name, buffer, 'application/pdf')
        }

        data = {
            'body': json.dumps({
                'username': report.user.username,
                'organization': report.organization.name_ru,
                'attack_type': report.attack_type
            })
        }
        try:
            url = 'https://cert.gov.kg/api/router/report-create'
            response = requests.post(url, data=data, files=files)
        except Exception as err:
            url = 'https://cert.gov.kg/api/router/report-create'
            buffer.seek(0)
            response = requests.post(url, data=data, files=files)

        buffer.seek(0)
        resp = FileResponse(buffer, as_attachment=True, filename=name, content_type='application/pdf')
        return resp


def get_static_analytic_file(request):

    reports = Report.objects.values('user__username', 'organization__name_en', 'risk_assessment').annotate(count=Count('id'))

    start = datetime.fromisoformat(request.GET.get('start'))
    end = datetime.fromisoformat(request.GET.get('end'))

    reports = reports.filter(detection_date__range=(start, end))

    reports_dicts = {}

    for report in reports:
        reports_dicts[report['user__username']] = reports_dicts.get(report['user__username'], {})
        reports_dicts[report['user__username']][report['organization__name_en']] = reports_dicts[report['user__username']].get(report['organization__name_en'], {})
        reports_dicts[report['user__username']][report['organization__name_en']][report['risk_assessment']] = reports_dicts[report['user__username']][report['organization__name_en']].get(report['risk_assessment'], {})
        reports_dicts[report['user__username']][report['organization__name_en']][report['risk_assessment']] = report['count']


    output = create_stat_report(reports_dicts, start=start, end=end)
    # Кодируем файл в base64
    excel_file_base64 = base64.b64encode(output.getvalue()).decode('utf-8')

    # Формируем JSON-ответ
    response_data = {
        "reports": reports_dicts,
        "file": {
            "filename": "report_kcokb.xlsx",
            "content": excel_file_base64,
            "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        }
    }

    return JsonResponse(response_data)


def get_departments(request):
    qs = list(Organization.objects.values_list('name_en', 'name_ru', 'pk'))
    return JsonResponse(qs, safe=False)


def analytics_view(request):
    return render(request, 'reports/analytics.html', {})


def get_filtered_qs(start, end):
    key = 'all_reports'
    if start and end:
        key = f'{start}+{end}'

    filter_qs = cache.get(key)

    if not filter_qs:
        print('Кэш не сработал')
        if not (start and end):
            bounds = Report.objects.aggregate(
                min_date=Min('detection_date'),
                max_date=Max('detection_date')
            )
            if not bounds['min_date']:
                return False

            start = bounds['min_date']
            end = bounds['max_date']

        filter_qs = Report.objects.filter(
            detection_date__range=(start, end)
        ).exclude(Q(country__isnull=True) | Q(country=""))
        cache.set(key, filter_qs, 5)
    return filter_qs


def get_attack_types_for_chart(request):
    contex = {}
    department = request.GET.get('department')
    start = request.GET.get('start')
    end = request.GET.get('end')

    filtered = get_filtered_qs(start, end)
    if not filtered:
        contex.update({'labels': [], 'data': []})
        return JsonResponse(contex)


    if department:
        filtered = filtered.filter(organization=department)

    # Группировка
    qs = (
        filtered
        .values('attack_type')
        .annotate(count=Count('id'))
        .order_by('-count')
    )

    max_attack_types = 10   # вытаскивам первые 10 наиболее встречающихся видов атак
    top = list(qs[:max_attack_types])
    others = list(qs[max_attack_types:])

    labels = [item['attack_type'] for item in top]
    data = [item['count'] for item in top]
    total = sum(data)

    if others:
        other_sum = sum(item['count'] for item in others)
        labels.append('Прочие')
        data.append(round(other_sum / total * 100, 2))

    contex.update({
        'labels': labels,
        'data': data,
        'total': total,
        'start': start,
        'end': end,
    })
    return JsonResponse(contex)


def get_risk_assessments_reports(request):
    contex = {}
    department = request.GET.get('department')
    start = request.GET.get('start')
    end = request.GET.get('end')

    # Получаем границы из БД, если их нет
    filtered = get_filtered_qs(start, end)
    if not filtered:
        contex.update({'labels': [], 'data': []})
        return JsonResponse(contex)

    if department:
        filtered = filtered.filter(organization=department)

    # Группировка
    qs = (
        filtered
        .values('risk_assessment')
        .annotate(count=Count('id'))
        .order_by('-count')
    )

    labels = [item['risk_assessment'] for item in qs]
    data = [item['count'] for item in qs]
    total = sum(data)

    contex.update({
        'labels': labels,
        'data': data,
        'total': total,
        'start': start,
        'end': end,
    })
    return JsonResponse(contex)


def get_countries_attacks(request):
    contex = {}
    department = request.GET.get('department')
    start = request.GET.get('start')
    end = request.GET.get('end')

    # Получаем границы из БД, если их нет
    filtered = get_filtered_qs(start, end)

    if not filtered:
        contex.update({'labels': [], 'data': []})
        return JsonResponse(contex)

    if department:
        filtered = filtered.filter(Q(organization=department) | Q(country__isnull=False))

    # Группировка
    qs = (
        filtered
        .values('country')
        .annotate(count=Count('id'))
        .order_by('-count')
    )

    max_attack_types = 10   # вытаскивам первые 10 наиболее встречающихся видов атак
    top = list(qs[:max_attack_types])
    others = list(qs[max_attack_types:])

    labels = [item['country'] for item in top]
    data = [item['count'] for item in top]
    total = sum(data)

    if others:
        other_sum = sum(item['count'] for item in others)
        labels.append('Прочие')
        data.append(other_sum)

    contex.update({
        'labels': labels,
        'data': data,
        'total': total,
        'start': start,
        'end': end,
    })
    return JsonResponse(contex)


def get_ip_count(request):
    contex = {}
    department = request.GET.get('department')
    start = request.GET.get('start')
    end = request.GET.get('end')

    # Получаем границы из БД, если их нет
    filtered = get_filtered_qs(start, end)

    if department:
        filtered = filtered.filter(organization=department)

    ip_lists = filtered.values('source_ip', 'country').annotate(count=Count('id')).order_by('-count')

    group_ip = {}
    for ip_obj in ip_lists[:10]:
        group_ip.setdefault(ip_obj['source_ip'], []).append([ip_obj['country'], ip_obj['count']])

    total = filtered.values('source_ip').annotate(count=Count('id')).count()
    contex.update({
        'data': group_ip,
        'total': total,
        'start': start,
        'end': end,
    })
    return JsonResponse(contex)


class ReportListView(ListView):
    model = Report
    template_name = 'reports/report_list.html'
    context_object_name = 'reports'
    paginate_by = 20

    def get_queryset(self):
        qs = super().get_queryset().order_by('-detection_date')
        start = self.request.GET.get('start')
        end = self.request.GET.get('end')
        if start and end:
            try:
                # Парсим даты
                sd_date = datetime.strptime(start, "%Y-%m-%d").date()
                ed_date = datetime.strptime(end, "%Y-%m-%d").date()
                # Границы суток
                sd = datetime.combine(sd_date, time.min)
                ed = datetime.combine(ed_date, time.max)
                qs = qs.filter(
                    detection_date__gte=sd,
                    detection_date__lte=ed
                )
            except ValueError:
                pass
        return qs

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['start'] = self.request.GET.get('start', '')
        ctx['end'] = self.request.GET.get('end', '')
        return ctx


class ReportDetailView(DetailView):
    model = Report
    template_name = 'reports/report_detail.html'
    context_object_name = 'report'


def export_monthly_reports(request):
    org_id = request.GET.get('org')
    start = request.GET.get('start')
    end = request.GET.get('end')

    rows = []
    if org_id and start and end:
        try:
            sd = datetime.strptime(start, "%Y-%m-%d")
            ed = datetime.strptime(end, "%Y-%m-%d") + timedelta(days=1)
        except ValueError:
            sd = ed = None

        if sd and ed:
            qs = Report.objects.filter(
                organization_id=org_id,
                detection_date__gte=sd,
                detection_date__lt=ed
            )
        else:
            qs = Report.objects.filter(organization_id=org_id)

        # Сбор данных по каждому source_ip…
        for src in qs.values_list('source_ip', flat=True).distinct():
            group = qs.filter(source_ip=src)
            cnt = group.count()
            rec = None
            for level, _ in RISK_LEVELS:
                rec = group.filter(risk_assessment=level).order_by('detection_date').first()
                if rec:
                    break
            rec = rec or group.order_by('detection_date').first()
            rows.append({
                'destination': ', '.join(group.values_list('destination_ip', flat=True).distinct()),
                'source': src,
                'warning': rec.attack_type,
                'priority': rec.risk_assessment,
            })

    # Генерация .docx…
    doc = Document()
    doc.add_heading(f'Ежемесячный отчет {start}–{end}', level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = '№'
    hdr[1].text = 'Destination IP'
    hdr[2].text = 'Source IP'
    hdr[3].text = 'Предупреждение'
    hdr[4].text = 'Критичность'

    for i, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = row['destination']
        cells[2].text = row['source']
        cells[3].text = row['warning']
        cells[4].text = row['priority']

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"monthly_report_{org_id}_{start}_{end}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response


class MonthlyReportView(TemplateView):
    template_name = 'reports/monthly_reports.html'

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['organizations'] = Organization.objects.order_by('name_en')
        ctx['risk_levels'] = RISK_LEVELS

        org_id = self.request.GET.get('org')
        start = self.request.GET.get('start')
        end = self.request.GET.get('end')
        ctx['selected_org'] = int(org_id) if org_id else None
        ctx['selected_start'] = start
        ctx['selected_end'] = end

        rows = []
        if org_id and start and end:
            try:
                sd_date = datetime.strptime(start, "%Y-%m-%d").date()
                ed_date = datetime.strptime(end, "%Y-%m-%d").date()
            except ValueError:
                ctx['reports'] = rows
                return ctx

            sd = datetime.combine(sd_date, time.min)
            ed = datetime.combine(ed_date, time.max)

            qs = Report.objects.filter(
                organization_id=org_id,
                detection_date__gte=sd,
                detection_date__lte=ed
            )

            for src in qs.values_list('source_ip', flat=True).distinct():
                group = qs.filter(source_ip=src)
                cnt = group.count()

                # Выбор записи по приоритету
                rec = None
                for level, _ in RISK_LEVELS:
                    rec = group.filter(risk_assessment=level) \
                        .order_by('detection_date') \
                        .first()
                    if rec:
                        break
                rec = rec or group.order_by('detection_date').first()

                rows.append({
                    'rec': rec,
                    'destination_ips': list(group.values_list('destination_ip', flat=True).distinct()),
                    'source_ip': src,
                    'warning': rec.attack_type,
                    'priority': rec.risk_assessment,
                    'requests': cnt,
                })

        ctx['reports'] = rows
        return ctx

    def post(self, request, *args, **kwargs):
        # Проходим по всем полям POST
        for key, val in request.POST.items():
            # Ищем только те, где имя начинается с manual_priority_
            if key.startswith('manual_priority_') and val:
                # Вычленяем pk записи
                rec_pk = key.split('manual_priority_')[1]
                # Обновляем только если пользователь выбрал значение (val != '')
                Report.objects.filter(pk=rec_pk).update(risk_assessment=val)

        # Сохраняем фильтры, чтобы после редиректа форма не сбросилась
        org = request.POST.get('org', '')
        start = request.POST.get('start', '')
        end = request.POST.get('end', '')
        return redirect(f"{request.path}?org={org}&start={start}&end={end}")
